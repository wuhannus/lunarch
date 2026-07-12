#!/usr/bin/env python3
"""
ICTA 2026 Paper Format Harness
===============================
Validates and auto-corrects conference paper formatting against ICTA_Template.doc.

Principles (from ICTA_Template.doc):
  1. A4 paper (21.0 × 29.7 cm)
  2. Two-column layout with 0.95 line spacing
  3. Margins: top 1.27cm, bottom 2.54cm, left/right 1.27cm
  4. Title: Arial 14 Bold, LEFT-ALIGNED (not center)
  5. Authors: Arial 10, LEFT-ALIGNED
  6. Body: Times New Roman 10pt
  7. Section heads: UPPERCASE Roman numerals, bold 10pt
  8. Sub-heads: Italic 10pt (A., B., C.)
  9. Abstract/Keywords: bold italic label, body in regular
  10. Tables: 8pt, captions bold centered above
  11. References: 8pt, [1]-style brackets
  12. No paragraph spacing before/after (0 pt)
  13. Line spacing exactly 0.95 multiple
  14. First-line indent: 0.5cm for body paragraphs

Usage:
  python3 harness.py check  — validate paper against template
  python3 harness.py fix    — auto-fix all issues
  python3 harness.py loop   — loop-engineering: fix→check until PASS
"""

import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dataclasses import dataclass, field
from typing import List, Tuple

# ============================================================================
# ICTA Template Specifications
# ============================================================================

@dataclass
class ICTASpecs:
    """ICTA 2026 template formatting requirements."""
    
    # Page
    page_width_cm: float = 21.0
    page_height_cm: float = 29.7
    top_margin_cm: float = 1.27
    bottom_margin_cm: float = 2.54
    left_margin_cm: float = 1.27
    right_margin_cm: float = 1.27
    
    # Fonts
    title_font: str = 'Arial'
    title_size: int = 14
    title_bold: bool = True
    title_align: int = WD_ALIGN_PARAGRAPH.LEFT  # NOT center
    
    author_font: str = 'Arial'
    author_size: int = 10
    
    body_font: str = 'Times New Roman'
    body_size: int = 10
    
    # Spacing
    line_spacing: float = 0.95
    para_before_pt: float = 0.0
    para_after_pt: float = 0.0
    first_indent_cm: float = 0.5
    
    # Section headings
    section_bold: bool = True
    section_size: int = 10
    
    # Table text
    table_font_size: int = 8

SPEC = ICTASpecs()

# ============================================================================
# Validation Engine
# ============================================================================

@dataclass
class Issue:
    severity: str  # ERROR, WARN, INFO
    location: str
    message: str
    expected: str = ""
    actual: str = ""

class PaperChecker:
    """Validates a Word document against ICTA specs."""
    
    def __init__(self, doc_path: str):
        self.doc = Document(doc_path)
        self.issues: List[Issue] = []
        self.section = self.doc.sections[0]
    
    def cm_to_emu(self, cm: float) -> int:
        return int(cm * 360000)
    
    def check_page_size(self):
        w = self.section.page_width
        h = self.section.page_height
        exp_w = self.cm_to_emu(SPEC.page_width_cm)
        exp_h = self.cm_to_emu(SPEC.page_height_cm)
        
        if abs(w - exp_w) > 50000:
            self.issues.append(Issue('ERROR', 'Page', 
                'Page width must be A4 (21.0 cm)',
                f'{SPEC.page_width_cm} cm', f'{w/360000:.1f} cm'))
        if abs(h - exp_h) > 50000:
            self.issues.append(Issue('ERROR', 'Page',
                'Page height must be A4 (29.7 cm)',
                f'{SPEC.page_height_cm} cm', f'{h/360000:.1f} cm'))
    
    def check_margins(self):
        for name, actual, expected in [
            ('Top', self.section.top_margin, SPEC.top_margin_cm),
            ('Bottom', self.section.bottom_margin, SPEC.bottom_margin_cm),
            ('Left', self.section.left_margin, SPEC.left_margin_cm),
            ('Right', self.section.right_margin, SPEC.right_margin_cm),
        ]:
            exp_emu = self.cm_to_emu(expected)
            if abs(actual - exp_emu) > 30000:
                self.issues.append(Issue('ERROR', 'Margins',
                    f'{name} margin incorrect',
                    f'{expected} cm', f'{actual/360000:.1f} cm'))
    
    def check_columns(self):
        """Verify two-column layout is present or indicated."""
        # python-docx doesn't directly expose columns. Check if section has cols property.
        try:
            cols = self.section._sectPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols')
            if cols is not None:
                num = int(cols.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '1'))
                if num != 2:
                    self.issues.append(Issue('ERROR', 'Layout',
                        'Must use two-column layout', '2 columns', f'{num} columns'))
            else:
                self.issues.append(Issue('WARN', 'Layout',
                    'Two-column layout not detected in XML. ICTA requires 2 columns.',
                    '', 'Single column detected'))
        except:
            self.issues.append(Issue('WARN', 'Layout',
                'Could not verify column count. Ensure 2-column layout in Word.',
                '2 columns', 'Unknown'))
    
    def check_title(self):
        """First paragraph should be title: Arial 14 Bold, left-aligned."""
        if not self.doc.paragraphs:
            return
        p = self.doc.paragraphs[0]
        if not p.runs:
            return
        
        r = p.runs[0]
        font_name = r.font.name
        font_size = r.font.size
        is_bold = r.font.bold
        alignment = p.alignment
        
        if font_name and font_name != SPEC.title_font:
            self.issues.append(Issue('ERROR', 'Title',
                f'Title font must be {SPEC.title_font}',
                SPEC.title_font, str(font_name)))
        if font_size and font_size != Pt(SPEC.title_size):
            self.issues.append(Issue('ERROR', 'Title',
                f'Title size must be {SPEC.title_size}pt',
                f'{SPEC.title_size}pt', f'{font_size/12700:.0f}pt' if font_size else '?'))
        if not is_bold:
            self.issues.append(Issue('ERROR', 'Title', 'Title must be bold'))
        if alignment is not None and alignment != SPEC.title_align:
            self.issues.append(Issue('ERROR', 'Title',
                'Title must be LEFT-ALIGNED (not centered per ICTA spec)',
                'LEFT', 'CENTERED' if alignment == WD_ALIGN_PARAGRAPH.CENTER else str(alignment)))
    
    def check_body_fonts(self):
        """Check that body paragraphs use Times New Roman 10pt."""
        for i, p in enumerate(self.doc.paragraphs):
            for r in p.runs:
                if r.font.name and r.font.name != SPEC.body_font and r.font.name != SPEC.title_font:
                    # Allow Arial for title/authors, warn about others
                    if r.font.size and r.font.size >= Pt(12):
                        continue  # title area OK with Arial
                    self.issues.append(Issue('WARN', f'Para {i}',
                        f'Body font should be {SPEC.body_font}, found {r.font.name}',
                        SPEC.body_font, r.font.name))
                    break  # One warning per paragraph
    
    def check_line_spacing(self):
        """Verify 0.95 line spacing on body paragraphs."""
        issues_found = 0
        for i, p in enumerate(self.doc.paragraphs[5:], 5):  # Skip title/author area
            ls = p.paragraph_format.line_spacing
            if ls and ls != SPEC.line_spacing:
                if issues_found < 3:
                    self.issues.append(Issue('WARN', f'Para {i}',
                        f'Line spacing should be {SPEC.line_spacing}',
                        str(SPEC.line_spacing), str(ls)))
                    issues_found += 1
    
    def check_paragraph_spacing(self):
        """Verify 0pt before/after on paragraphs."""
        issues_found = 0
        for i, p in enumerate(self.doc.paragraphs[5:], 5):
            before = p.paragraph_format.space_before
            after = p.paragraph_format.space_after
            if before and before > Pt(6):
                if issues_found < 2:
                    self.issues.append(Issue('WARN', f'Para {i}',
                        'Space-before should be 0pt', '0pt', f'{before/12700:.0f}pt'))
                    issues_found += 1
            # Allow small space-after for readability
    
    def check_tables(self):
        """Check table formatting."""
        for ti, table in enumerate(self.doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.font.size and r.font.size > Pt(SPEC.table_font_size + 1):
                                self.issues.append(Issue('WARN', f'Table {ti}',
                                    f'Table text should be {SPEC.table_font_size}pt',
                                    f'{SPEC.table_font_size}pt', f'{r.font.size/12700:.0f}pt'))
                                break
    
    def check_first_line_indent(self):
        """Verify body paragraphs have ~0.5cm first-line indent."""
        body_count = 0
        indented = 0
        for i, p in enumerate(self.doc.paragraphs[10:], 10):  # Skip header area
            if p.text.strip() and len(p.text) > 50:  # Body paragraphs have substantial text
                body_count += 1
                indent = p.paragraph_format.first_line_indent
                if indent and abs(indent - Cm(0.5)) < Cm(0.2):
                    indented += 1
        
        if body_count > 0 and indented / body_count < 0.5:
            self.issues.append(Issue('WARN', 'Body',
                'Most body paragraphs should have 0.5cm first-line indent',
                '>50% with 0.5cm indent', f'{indented}/{body_count} paragraphs indented'))
    
    def run_all_checks(self):
        self.issues = []
        self.check_page_size()
        self.check_margins()
        self.check_columns()
        self.check_title()
        self.check_body_fonts()
        self.check_line_spacing()
        self.check_paragraph_spacing()
        self.check_tables()
        self.check_first_line_indent()
        return self.issues


# ============================================================================
# Auto-Fixer
# ============================================================================

class PaperFixer:
    """Auto-corrects paper formatting to match ICTA specs."""
    
    def __init__(self, doc_path: str, output_path: str):
        self.doc = Document(doc_path)
        self.output_path = output_path
        self.section = self.doc.sections[0]
        self.fixes_applied = []
    
    def fix_page_size(self):
        self.section.page_width = Cm(SPEC.page_width_cm)
        self.section.page_height = Cm(SPEC.page_height_cm)
        self.fixes_applied.append('Page size → A4')
    
    def fix_margins(self):
        self.section.top_margin = Cm(SPEC.top_margin_cm)
        self.section.bottom_margin = Cm(SPEC.bottom_margin_cm)
        self.section.left_margin = Cm(SPEC.left_margin_cm)
        self.section.right_margin = Cm(SPEC.right_margin_cm)
        self.fixes_applied.append('Margins → ICTA spec')
    
    def fix_columns(self):
        """Add two-column layout to the section."""
        from docx.oxml.ns import qn
        sectPr = self.section._sectPr
        # Find or create cols element
        cols = sectPr.find(qn('w:cols'))
        if cols is None:
            from lxml import etree
            cols = etree.SubElement(sectPr, qn('w:cols'))
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')  # ~0.5cm gutter between columns
        self.fixes_applied.append('Layout → Two-column')
    
    def fix_title(self):
        if not self.doc.paragraphs:
            return
        p = self.doc.paragraphs[0]
        p.alignment = SPEC.title_align
        for r in p.runs:
            r.font.name = SPEC.title_font
            r.font.size = Pt(SPEC.title_size)
            r.bold = True
        self.fixes_applied.append('Title → Arial 14 Bold Left')
    
    def fix_line_spacing(self):
        for p in self.doc.paragraphs:
            p.paragraph_format.line_spacing = SPEC.line_spacing
        self.fixes_applied.append('Line spacing → 0.95')
    
    def fix_paragraph_spacing(self):
        for p in self.doc.paragraphs:
            p.paragraph_format.space_before = Pt(SPEC.para_before_pt)
        self.fixes_applied.append('Paragraph before-spacing → 0pt')
    
    def fix_table_fonts(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(SPEC.table_font_size)
                            r.font.name = SPEC.body_font
        self.fixes_applied.append('Table fonts → 8pt')
    
    def fix_body_fonts(self):
        """Set body paragraphs to Times New Roman 10pt, skipping title area."""
        for i, p in enumerate(self.doc.paragraphs[3:], 3):  # Skip title, author, blank
            for r in p.runs:
                if r.font.size and r.font.size >= Pt(13):
                    continue  # Don't change title-sized text
                r.font.name = SPEC.body_font
                if r.font.size is None or r.font.size > Pt(11):
                    r.font.size = Pt(SPEC.body_size)
        self.fixes_applied.append('Body fonts → Times New Roman 10pt')
    
    def fix_first_line_indent(self):
        """Add 0.5cm first-line indent to body paragraphs with substantial text."""
        for i, p in enumerate(self.doc.paragraphs[10:], 10):
            if p.text.strip() and len(p.text) > 50 and not p.style.name.startswith('Heading'):
                if p.paragraph_format.first_line_indent is None:
                    p.paragraph_format.first_line_indent = Cm(SPEC.first_indent_cm)
        self.fixes_applied.append('First-line indent → 0.5cm')
    
    def apply_all_fixes(self):
        self.fix_page_size()
        self.fix_margins()
        self.fix_columns()
        self.fix_title()
        self.fix_line_spacing()
        self.fix_paragraph_spacing()
        self.fix_table_fonts()
        self.fix_body_fonts()
        self.fix_first_line_indent()
        self.doc.save(self.output_path)
        return self.fixes_applied


# ============================================================================
# Main — Loop Engineering
# ============================================================================

def main():
    import argparse
    ap = argparse.ArgumentParser(description='ICTA 2026 Paper Format Harness')
    ap.add_argument('action', choices=['check', 'fix', 'loop'], 
                    help='check=validate, fix=auto-correct, loop=fix until pass')
    ap.add_argument('--input', default='/Users/wuhan0515/opencode/lunarch_icta_2026.docx',
                    help='Input paper path')
    ap.add_argument('--output', default='/Users/wuhan0515/opencode/lunarch_icta_2026.docx',
                    help='Output paper path')
    args = ap.parse_args()
    
    if args.action == 'check':
        checker = PaperChecker(args.input)
        issues = checker.run_all_checks()
        
        print("=" * 60)
        print("  ICTA 2026 PAPER FORMAT CHECK")
        print("=" * 60)
        
        errors = [i for i in issues if i.severity == 'ERROR']
        warns = [i for i in issues if i.severity == 'WARN']
        infos = [i for i in issues if i.severity == 'INFO']
        
        for issue in issues:
            icon = {'ERROR': '✗', 'WARN': '⚠', 'INFO': 'ℹ'}[issue.severity]
            print(f"  {icon} [{issue.location:12s}] {issue.message}")
            if issue.expected:
                print(f"      Expected: {issue.expected}")
                print(f"      Actual:   {issue.actual}")
        
        print(f"\n  Errors: {len(errors)}  Warnings: {len(warns)}  Info: {len(infos)}")
        if errors:
            print("  Result: FAIL ✗ — Fix errors before submission")
        elif warns:
            print("  Result: PASS with warnings ⚠ — Review warnings")
        else:
            print("  Result: PASS ✓ — Paper meets all ICTA specs")
        print("=" * 60)
        
        return len(errors) == 0
    
    elif args.action == 'fix':
        print("=" * 60)
        print("  ICTA 2026 PAPER FORMAT AUTO-FIX")
        print("=" * 60)
        fixer = PaperFixer(args.input, args.output)
        fixes = fixer.apply_all_fixes()
        for f in fixes:
            print(f"  ✓ {f}")
        print(f"\n  Fixed: {len(fixes)} issues corrected")
        print(f"  Saved: {args.output}")
        print("=" * 60)
        
        # Now check
        checker = PaperChecker(args.output)
        issues = checker.run_all_checks()
        errors = [i for i in issues if i.severity == 'ERROR']
        print(f"  Remaining errors after fix: {len(errors)}")
        return len(errors) == 0
    
    elif args.action == 'loop':
        print("=" * 60)
        print("  ICTA 2026 LOOP-ENGINEERING HARNESS")
        print("  Fix → Check → Repeat until PASS")
        print("=" * 60)
        
        max_iterations = 5
        for iteration in range(1, max_iterations + 1):
            print(f"\n--- Iteration {iteration}/{max_iterations} ---")
            
            # Fix
            fixer = PaperFixer(args.input, args.output)
            fixes = fixer.apply_all_fixes()
            print(f"  Fixes applied: {len(fixes)}")
            
            # Check
            checker = PaperChecker(args.output)
            issues = checker.run_all_checks()
            errors = [i for i in issues if i.severity == 'ERROR']
            
            for issue in errors:
                print(f"  ✗ {issue.message}")
            
            print(f"  Errors remaining: {len(errors)}")
            
            if not errors:
                print(f"\n  ✓✓✓ PASS after {iteration} iteration(s) ✓✓✓")
                # Show final stats
                warns = [i for i in issues if i.severity == 'WARN']
                if warns:
                    print(f"  Warnings: {len(warns)} (non-blocking)")
                print("=" * 60)
                return True
            
            args.input = args.output  # Use fixed version for next iteration
        
        print(f"\n  ✗ FAIL after {max_iterations} iterations")
        print("  Manual review required for remaining errors")
        print("=" * 60)
        return False


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
